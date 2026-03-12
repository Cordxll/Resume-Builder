"""DOCX export functionality to create tailored resume documents"""
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_to_docx(
    resume: dict,
    tailored_data: dict,
    accepted_changes: dict,
    output_path: str
) -> None:
    """
    Export resume to DOCX format with accepted changes applied.

    Args:
        resume:          Full ParsedResume dict — contact + sections[].
        tailored_data:   TailoredData dict — summary/skills objects, jobs[], sections[].
        accepted_changes: Per-section booleans, e.g. {"summary": True, "experience": True, ...}
        output_path:     Absolute path where the .docx file will be saved.
    """
    doc = Document()

    # Set document margins
    for sec in doc.sections:
        sec.top_margin = Inches(0.5)
        sec.bottom_margin = Inches(0.5)
        sec.left_margin = Inches(0.75)
        sec.right_margin = Inches(0.75)

    # ── Contact ──────────────────────────────────────────────────────────────
    contact = resume.get("contact") or {}
    add_contact_section(doc, contact)

    # Build a lookup of tailored jobs by id: { "job-0": TailoredJobEntry, ... }
    tailored_jobs_by_id: Dict[str, dict] = {
        job["id"]: job
        for job in tailored_data.get("jobs", [])
        if job.get("id")
    }

    # Build a lookup of TailoredSection by sectionType:
    # { "projects": TailoredSection, "education": TailoredSection, ... }
    tailored_sections_by_type: Dict[str, dict] = {
        ts["sectionType"]: ts
        for ts in tailored_data.get("sections", [])
        if ts.get("sectionType")
    }

    # ── Sections (iterate in original order) ─────────────────────────────────
    for section in resume.get("sections", []):
        section_type: str = section.get("type", "unknown")
        raw_label: str = section.get("rawLabel", section_type.capitalize())
        accepted: bool = accepted_changes.get(section_type, True)

        add_section_header(doc, raw_label.upper())

        if section_type == "summary":
            _write_summary(doc, section, tailored_data, accepted)

        elif section_type == "skills":
            _write_skills(doc, section, tailored_data, accepted)

        elif section_type == "experience":
            _write_experience(doc, section, tailored_jobs_by_id, accepted)

        elif section_type == "education":
            _write_education(doc, section, tailored_sections_by_type, accepted)

        elif section_type in ("projects", "certifications", "awards", "volunteer"):
            _write_generic_entries(
                doc, section, tailored_sections_by_type, accepted, section_type
            )

        else:
            # languages, contact (duplicate), unknown, or any future type
            _write_fallback(doc, section)

    doc.save(output_path)


# ── Section writers ────────────────────────────────────────────────────────────

def _write_summary(
    doc: Document,
    section: dict,
    tailored_data: dict,
    accepted: bool,
) -> None:
    if accepted:
        text = (tailored_data.get("summary") or {}).get("tailored", "") or section.get("content", "")
    else:
        text = section.get("content", "")

    if text and text.strip():
        p = doc.add_paragraph(text.strip())
        set_paragraph_font(p, size=10.5)
    doc.add_paragraph()


def _write_skills(
    doc: Document,
    section: dict,
    tailored_data: dict,
    accepted: bool,
) -> None:
    if accepted:
        text = (tailored_data.get("skills") or {}).get("tailored", "") or section.get("content", "")
    else:
        text = section.get("content", "")

    if text and text.strip():
        p = doc.add_paragraph(text.strip())
        set_paragraph_font(p, size=10.5)
    doc.add_paragraph()


def _write_experience(
    doc: Document,
    section: dict,
    tailored_jobs_by_id: Dict[str, dict],
    accepted: bool,
) -> None:
    jobs: List[dict] = section.get("jobs", [])
    for i, job in enumerate(jobs):
        job_id: str = job.get("id", "")

        # ── Header row: Company + dates/location ─────────────────────────────
        p = doc.add_paragraph()
        company_run = p.add_run(job.get("company", ""))
        company_run.font.size = Pt(11)
        company_run.font.bold = True
        company_run.font.color.rgb = RGBColor(0, 0, 0)

        dates = job.get("dates", "")
        location = job.get("location", "")
        if dates and location:
            right_text = f"  |  {dates}  |  {location}"
        elif dates:
            right_text = f"  |  {dates}"
        elif location:
            right_text = f"  |  {location}"
        else:
            right_text = ""

        if right_text:
            date_run = p.add_run(right_text)
            date_run.font.size = Pt(10)
            date_run.font.color.rgb = RGBColor(100, 100, 100)

        # ── Title ─────────────────────────────────────────────────────────────
        title = job.get("title", "")
        if title:
            title_p = doc.add_paragraph()
            title_run = title_p.add_run(title)
            title_run.font.size = Pt(10.5)
            title_run.font.italic = True
            title_run.font.color.rgb = RGBColor(60, 60, 60)

        # ── Bullets ───────────────────────────────────────────────────────────
        if accepted and job_id in tailored_jobs_by_id:
            tailored_job = tailored_jobs_by_id[job_id]
            # tailored bullets are TailoredBullet: {index, original, tailored, explanation}
            bullets: List[dict] = tailored_job.get("bullets", [])
            bullet_texts = [b.get("tailored", b.get("original", "")) for b in bullets]
        else:
            # original bullets are Bullet: {text, index}
            raw_bullets: List[dict] = job.get("bullets", [])
            bullet_texts = [b.get("text", "") for b in raw_bullets]

        for text in bullet_texts:
            if text and str(text).strip():
                bp = doc.add_paragraph(str(text).strip(), style="List Bullet")
                set_paragraph_font(bp, size=10.5)

        # Spacing between jobs
        if i < len(jobs) - 1:
            doc.add_paragraph()

    doc.add_paragraph()


def _write_education(
    doc: Document,
    section: dict,
    tailored_sections_by_type: Dict[str, dict],
    accepted: bool,
) -> None:
    entries: List[dict] = section.get("entries", [])
    tailored_section: Optional[dict] = tailored_sections_by_type.get("education")

    # Build tailored bullet lookup: { entry_id: [TailoredBullet, ...] }
    tailored_bullets_by_id: Dict[str, List[dict]] = {}
    if tailored_section:
        for entry in tailored_section.get("entries", []):
            eid = entry.get("id", "")
            if eid:
                tailored_bullets_by_id[eid] = entry.get("bullets", [])

    for i, entry in enumerate(entries):
        entry_id = entry.get("id", "")

        # Header: institution + dates
        p = doc.add_paragraph()
        inst_run = p.add_run(entry.get("institution", ""))
        inst_run.font.size = Pt(11)
        inst_run.font.bold = True
        inst_run.font.color.rgb = RGBColor(0, 0, 0)

        dates = entry.get("dates", "")
        if dates:
            date_run = p.add_run(f"  |  {dates}")
            date_run.font.size = Pt(10)
            date_run.font.color.rgb = RGBColor(100, 100, 100)

        # Degree / field line
        degree = entry.get("degree", "")
        field = entry.get("field", "")
        degree_line = ", ".join(filter(None, [degree, field]))
        if degree_line:
            deg_p = doc.add_paragraph()
            deg_run = deg_p.add_run(degree_line)
            deg_run.font.size = Pt(10.5)
            deg_run.font.italic = True
            deg_run.font.color.rgb = RGBColor(60, 60, 60)

        # Bullets
        if accepted and entry_id in tailored_bullets_by_id:
            bullet_texts = [
                b.get("tailored", b.get("original", ""))
                for b in tailored_bullets_by_id[entry_id]
            ]
        else:
            raw_bullets: List[dict] = entry.get("bullets", [])
            bullet_texts = [b.get("text", "") for b in raw_bullets]

        for text in bullet_texts:
            if text and str(text).strip():
                bp = doc.add_paragraph(str(text).strip(), style="List Bullet")
                set_paragraph_font(bp, size=10.5)

        if i < len(entries) - 1:
            doc.add_paragraph()

    doc.add_paragraph()


def _write_generic_entries(
    doc: Document,
    section: dict,
    tailored_sections_by_type: Dict[str, dict],
    accepted: bool,
    section_type: str,
) -> None:
    """
    Handle projects, certifications, awards, volunteer, and any future entry-based
    section types using the same TailoredSection lookup pattern.
    """
    entries: List[dict] = section.get("entries", [])
    tailored_section: Optional[dict] = tailored_sections_by_type.get(section_type)

    # Build tailored bullet lookup: { entry_id: [TailoredBullet, ...] }
    tailored_bullets_by_id: Dict[str, List[dict]] = {}
    if tailored_section:
        for entry in tailored_section.get("entries", []):
            eid = entry.get("id", "")
            if eid:
                tailored_bullets_by_id[eid] = entry.get("bullets", [])

    for i, entry in enumerate(entries):
        entry_id = entry.get("id", "")

        # ── Entry header line ─────────────────────────────────────────────────
        # Use 'name' for projects/certs/awards; fall back to generic fields
        name = (
            entry.get("name")
            or entry.get("institution")
            or entry.get("content", "")
        )
        p = doc.add_paragraph()
        name_run = p.add_run(name or "")
        name_run.font.size = Pt(11)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0, 0, 0)

        # Supplementary info: issuer/date for certs & awards; dates for projects
        issuer = entry.get("issuer", "")
        date = entry.get("date", "") or entry.get("dates", "")
        meta_parts = [p for p in [issuer, date] if p]
        if meta_parts:
            meta_run = p.add_run(f"  |  {' | '.join(meta_parts)}")
            meta_run.font.size = Pt(10)
            meta_run.font.color.rgb = RGBColor(100, 100, 100)

        # Optional description line (projects may have one)
        description = entry.get("description", "")
        if description:
            desc_p = doc.add_paragraph()
            desc_run = desc_p.add_run(description)
            desc_run.font.size = Pt(10.5)
            desc_run.font.italic = True
            desc_run.font.color.rgb = RGBColor(60, 60, 60)

        # ── Bullets ───────────────────────────────────────────────────────────
        if accepted and entry_id in tailored_bullets_by_id:
            bullet_texts = [
                b.get("tailored", b.get("original", ""))
                for b in tailored_bullets_by_id[entry_id]
            ]
        else:
            raw_bullets: List[dict] = entry.get("bullets", [])
            bullet_texts = [b.get("text", "") for b in raw_bullets]

        for text in bullet_texts:
            if text and str(text).strip():
                bp = doc.add_paragraph(str(text).strip(), style="List Bullet")
                set_paragraph_font(bp, size=10.5)

        if i < len(entries) - 1:
            doc.add_paragraph()

    doc.add_paragraph()


def _write_fallback(doc: Document, section: dict) -> None:
    """Write content for languages, duplicate contact, unknown, or generic sections."""
    content = section.get("content", "")
    if content and content.strip():
        p = doc.add_paragraph(content.strip())
        set_paragraph_font(p, size=10.5)
    else:
        # Try entries with optional content/name fields
        for entry in section.get("entries", []):
            line = entry.get("content") or entry.get("name") or ""
            if line and line.strip():
                p = doc.add_paragraph(line.strip())
                set_paragraph_font(p, size=10.5)
    doc.add_paragraph()


# ── Shared helpers ─────────────────────────────────────────────────────────────

def add_contact_section(doc: Document, contact: dict) -> None:
    """Add contact information block to the document."""
    if not contact:
        return

    # Name as centred title
    name = contact.get("name", "")
    if name:
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(name)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

    # One-line contact details
    contact_parts = []
    for key in ("email", "phone", "linkedin", "github", "website"):
        value = contact.get(key, "")
        if value:
            contact_parts.append(value)

    if contact_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(" | ".join(contact_parts))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()


def add_section_header(doc: Document, title: str) -> None:
    """Add a bold section heading followed by a thin horizontal rule."""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    rule = doc.add_paragraph()
    run = rule.add_run("_" * 85)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(150, 150, 150)


def set_paragraph_font(paragraph, size: float = 10.5) -> None:
    """Set font size for all runs in a paragraph."""
    for run in paragraph.runs:
        run.font.size = Pt(size)
